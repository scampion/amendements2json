import difflib
import json
import sys
import logging
import datetime

from bs4 import BeautifulSoup
import re

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
import requests
from io import BytesIO

import am2json.meps as meps
from am2json.am_labeler import get_label_am
from am2json.dossier_retriever import get_final_dossier_link

logging.basicConfig(level=logging.DEBUG)
log = logging.getLogger(__name__)

mep_info = meps.get_mep_data()

legislatures_dates = {7: (2009, 2014),
                      8: (2014, 2019),
                      9: (2019, 2024)}


def get_legislature(date):
    date = datetime.datetime.strptime(date, '%Y-%m-%d').date()
    for legislature, dates in legislatures_dates.items():
        if datetime.date(dates[0], 7, 1) <= date <= datetime.date(dates[1], 7, 1):
            return legislature


def clean(func):
    def wrapper(*args, **kwargs):
        text = func(*args, **kwargs)
        for r in re.findall(r'\{(.*?)\}', text):
            text = r
            break
        for r in re.findall(r'\((.*?)\)$', text):
            text = r
            break
        return text

    return wrapper


# def get_html(doc_file):
#     doc = Document(doc_file)
#     text_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
#     soup = BeautifulSoup(text_content, 'html.parser')
#     return soup


def get_html(file):
    def iter_block_items(parent):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        elif isinstance(parent, _Row):
            parent_elm = parent._tr
        else:
            raise ValueError("something's not right")
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    document = Document(file)

    html_string = ""
    for block in iter_block_items(document):
        # print(block.text if isinstance(block, Paragraph) else '<table>')
        if isinstance(block, Paragraph):
            if block.text:
                html_string += block.text + "<br>"
        elif isinstance(block, Table):
            html_string += "<table>"
            for row in block.rows:
                if any(cell.text.strip() for cell in row.cells):  # Check if any cell in the row is non-empty
                    html_string += "<tr>"
                    row_data = []
                    for cell in row.cells:
                        cell_text = " ".join([paragraph.text for paragraph in cell.paragraphs])
                        html_string += f"<td>{cell_text}</td>"
                    html_string += "</tr>"

            html_string += "</table>"

    soup = BeautifulSoup(html_string, "html.parser")
    return soup


# Parameter local defines whether to get the file from local directory
# Otherwise get it online


# Based on docx file of final amendment, get all amendments
def extract_final_amendments(file):

    document = Document(file)

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    amendments = []
    # Minimum paragraph length
    min_length = 12
    for paragraph in paragraphs:
        # Check if the paragraph starts with an enumeration pattern, i.e. with "-" or "A." or 1.
        if re.match(r'^[-\dA-Z]+\.', paragraph) or paragraph[0] in ["-", "–"]:
            # Check if tab character is somewhere in the first few chartacters it to make sure it's an numeration
            # the len(paragraph) > 12 is to avoid bad edge cases like " - : against" key at the end of document
            if len(paragraph) > min_length and '\t' in paragraph[0:8]:
                # skip tab when adding text
                offset_index = paragraph.find('\t')
                text = paragraph[(offset_index + 1):]
                amendments.append(text)
    return amendments


# Based on dossier ID, look for link or local file to get final amendments docx, and then apply function above.
def get_final_dossier_am(dossier_id, local=False):

    if local:
        dossier_file = f"./final_dossiers/{dossier_id}.docx"
        amendments = extract_final_amendments(dossier_file)
        return amendments
    else:
        url_link = get_final_dossier_link(dossier_id)
        if url_link:
            response = requests.get(url_link)
            file_in_mem = BytesIO(response.content)
            amendments = extract_final_amendments(file_in_mem)
            return amendments
        else:
            return []


# directory has to be a Pathlib directory, gets all final dossiers based on
# a directory that contains the amendment documents.
def get_final_dossiers(directory, download=False, debug=False):
    doc_files = list(directory.rglob("*_EN.docx"))
    final_dossiers = []

    for doc_file in doc_files:
        soup = get_html(doc_file)
        dossier_id = get_dossier_id(soup)
        final_dossier = get_final_dossier_link(dossier_id, download=download, debug=debug)
        final_dossiers.append(final_dossier)
    return final_dossiers


def get_parliament(soup):
    if soup.find("td"):
        return soup.find("td").text.strip()
    else:
        log.warning("No parliament found")


@clean
def get_committee(soup):
    return soup.find("commission").text.strip().replace("Committee on ", "")


@clean
def get_dossier_ref(soup):
    if soup.find("refproc"):
        return soup.find("refproc").text.strip()
    elif soup.find("docrefpe"):
        return soup.find("docrefpe").text.strip()


@clean
def get_source(soup):
    if soup.find("docref"):
        return soup.find("docref").text.strip()
    elif soup.find("docrefpe"):
        return soup.find("docrefpe").text.strip()


@clean
def get_dossier_id(soup):
    """
    The dossier ID, i.e. PE700.718, are always in docrefpre tags.
    """
    try:
        dossier_id = soup.find("docrefpe").text.strip()
        dossier_id = dossier_id.strip('()')
        # Get only
        dossier_id = dossier_id[:9]
        return dossier_id
    except:
        return "None"


@clean
def get_date(soup):
    @clean
    def get_raw_date(soup):
        return soup.find("date").text.strip()

    date = datetime.datetime.strptime(get_raw_date(soup), '%d/%m/%Y').date()
    return date.strftime('%Y-%m-%d')


def get_article_type(soup):
    if soup.find("article"):
        article_type = soup.find("article").text.strip()
        return article_type.split(' ')[0].lower()


def get_titretype(soup):
    if soup.find('titretype'):
        return soup.find('titretype').text.strip().split()[-1].lower()


def get_legal_act(soup):
    # Find the <docref> tag cause the legal type is after the <br> right next to it.
    try:
        docref_tag = soup.find('docref')
        br_tag = docref_tag.find_next('br')
        leg_type = br_tag.next_sibling.string.strip()
    except:  # TODO Too broad exception clause but I don't know what else to do.
        leg_type = "NA"
    # Check if document is a resolution
    if "Non" in leg_type:
        return "resolution"
    elif soup.find(text="Proposal for a directive"):
        return "directive"
    elif soup.find(text="Proposal for a regulation"):
        return "regulation"
    elif soup.find(text="Proposal for a decision"):
        return "decision"
    else:
        return "None"


def get_metadata(soup):
    return {'committee': get_committee(soup),
            'dossier_ref': get_dossier_ref(soup),
            'dossier_id': get_dossier_id(soup),
            'date': get_date(soup),
            'rapporteur': soup.find("rapporteur").text.strip(),
            'source': get_source(soup),
            'dossier_title': soup.find("titre").text.strip(),
            'dossier_type': get_titretype(soup),
            'legal_act': get_legal_act(soup),
            'justification': None
            }


def get_amd(amend):
    table = amend.find("table")
    text_original = table.find_all("tr")[1].find_all("td")[0].text.strip()
    text_amended = table.find_all("tr")[1].find_all("td")[1].text.strip()

    # We test if the first word match a regex to detect only one letter between parenthesis
    # If it's the case, we remove it
    if text_original and re.match(r'\([a-zA-Z]\)', text_original.split()[0]):
        text_original = ' '.join(text_original.split()[1:])
    if text_amended and re.match(r'\([a-zA-Z]\)', text_amended.split()[0]):
        text_amended = ' '.join(text_amended.split()[1:])

    text_original = text_original.replace(',', ' ,')
    text_amended = text_amended.replace(',', ' ,')

    # We get here the edit_indices and edit type using difflib
    diff = difflib.ndiff(text_original.split(), text_amended.split())
    edit_type = None
    diff_indices = {'i1': 0, 'i2': 0, 'j1': 0, 'j2': 0}

    for i, line in enumerate(diff):
        if line.startswith('-'):
            edit_type = 'delete'
            diff_indices['i1'] = i
            diff_indices['i2'] = i + 1
            diff_indices['j1'] = i
            diff_indices['j2'] = i
            break
        elif line.startswith('+'):
            edit_type = 'insert'
            diff_indices['i1'] = i
            diff_indices['i2'] = i
            diff_indices['j1'] = i
            diff_indices['j2'] = i + 1
            break
        elif line.startswith('?'):
            edit_type = 'replace'
            diff_indices['i1'] = i
            diff_indices['i2'] = i + 1
            diff_indices['j1'] = i
            diff_indices['j2'] = i + 1
            break
    return text_original.lower().split(), text_amended.lower().split(), edit_type, diff_indices


def get_authors(amend, date):
    def get_mep_info(name):
        name = name.replace('ß', 'ss')  # Albert DESS
        name = name.replace('–', '-')  # "jean–paul denanot" != "jean-paul denanot"
        global mep_info
        for mep_id, mep_data in mep_info.items():
            if mep_data.get('name').lower() == name.lower():
                return mep_id, mep_data
        log.warning(f"Could not find MEP {name} in the list of MEPs")

    for author in [author.strip() for author in amend.find("members").text.split(",")]:
        mep_id, mep_data = get_mep_info(author)
        if mep_data:
            yield {
                'id': int(mep_id),
                'name': mep_data.get('name'),
                'gender': mep_data.get('gender'),
                'nationality': mep_data.get('nationality'),
                'group': mep_data.get(f'group-ep{get_legislature(date)}'),
                'rapporteur': False
            }


def get_amend_num(amend):
    if amend.find("numam"):
        return amend.find("numam").text.strip()


def get_justification(amend):
    if amend.find("titrejust"):
        titrejust = amend.find("titrejust")
        # find the next child of this node titrejust
        # that is not a <br> tag
        # and return its text
        return titrejust.find_next_sibling(text=True).strip()


def get_amendments(soup):
    md = get_metadata(soup)

    # Extract html from final amendment
    dossier_id = md["dossier_id"]
    final_amendments= get_final_dossier_am(dossier_id)
    for amend in soup.find_all("amend"):
        md['article_type'] = get_article_type(amend)
        md['target'] = soup.find('article').text.strip()
        md['authors'] = list(get_authors(amend, md['date']))
        md['amendment_num'] = get_amend_num(amend)
        md['justification'] = get_justification(amend)
        md['text_original'], md['text_amended'], md['edit_type'], md['edit_indices'] = get_amd(amend)
        md["accepted"] = get_label_am(md['text_amended'], final_amendments)
        yield md

def extract_amendments(file):
    soup = get_html(file)
    assert soup.find("typeam") and soup.find("typeam").text.strip() == "AMENDMENTS", "Not an amendment file"
    md = get_metadata(soup)
    for i, a in enumerate(get_amendments(soup)):
        yield md | a


if __name__ == '__main__':
    soup = get_html(sys.argv[1])
    with open(sys.argv[1] + ".json", "w") as f:
        json.dump(soup.prettify(), f, indent=2)
    for i, r in enumerate(extract_amendments(sys.argv[1])):
        with open(sys.argv[1] + f"_{i}.json", 'w') as f:
            json.dump(r, f, indent=2, separators=(',', ':'))
        print(json.dumps(r, indent=2, separators=(',', ':')))
        print("-" * 80)
